import base64
import io
import json
import logging
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import fitz
from PIL import Image
import anthropic
from langchain.docstore.document import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dataclasses import dataclass
from docx import Document as DocxDocument
import tempfile
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit

from .config import settings

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Structure d'un chunk de document avec métadonnées enrichies"""
    content: str
    metadata: Dict
    bbox: Tuple[int, int, int, int]
    confidence: float

class VisualDocumentAnalyzer:
    """Analyseur de documents basé sur Claude Vision pour l'extraction de texte médical"""
    def __init__(self, anthropic_api_key: Optional[str] = None):
        api_key = anthropic_api_key or settings.anthropic_api_key
        if not api_key:
            raise ValueError("Clé API Anthropic requise")
            
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-3-haiku-20240307"
        
    def convert_doc_to_images(self, doc_path: str, dpi: int = 200) -> List[Image.Image]:
        """Convertit un document (PDF/DOCX) en images haute résolution"""
        
        # Si c'est un .docx, le convertir en PDF d'abord
        if doc_path.endswith('.docx'):
            pdf_path = self._docx_to_pdf(doc_path)
        else:
            pdf_path = doc_path
            
        # Extraire les images du PDF
        doc = fitz.open(pdf_path)
        images = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # Matrice de transformation pour haute résolution
            mat = fitz.Matrix(dpi/72, dpi/72)
            pix = page.get_pixmap(matrix=mat)
            
            # Convertir en PIL
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))
            images.append(img)  
            
        doc.close()
        return images
    
    def _docx_to_pdf(self, docx_path: str) -> str:
        """Convertit un fichier DOCX en PDF temporaire"""
        return self._extract_text_from_docx(docx_path)
    
    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extrait le texte d'un DOCX et génère un PDF temporaire"""
        try:
            doc = DocxDocument(docx_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    if paragraph.style.name.startswith('Heading'):
                        text_content.append(f"\n*** {text} ***\n")
                    else:
                        text_content.append(text)
            for table in doc.tables:
                text_content.append("\n=== TABLEAU ===")
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_content.append(row_text)
                text_content.append("=== FIN TABLEAU ===\n")
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                pdf_path = tmp.name
                
            c = canvas.Canvas(pdf_path, pagesize=letter)
            width, height = letter
            c.setFont("Helvetica", 10)
            y_position = height - 50
            margin = 50
            max_width = width - 2 * margin
            
            for paragraph in text_content:
                if not paragraph.strip():
                    continue
                    
                if paragraph.startswith("***") and paragraph.endswith("***"):
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 50
                    c.setFont("Helvetica-Bold", 12)
                    c.drawString(margin, y_position, paragraph.replace("*", "").strip())
                    y_position -= 20
                    c.setFont("Helvetica", 10)
                    continue
                
                if paragraph.startswith("==="):
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    c.setFont("Helvetica-Bold", 9)
                    c.drawString(margin, y_position, paragraph)
                    y_position -= 15
                    c.setFont("Helvetica", 10)
                    continue
                
                lines = simpleSplit(paragraph, "Helvetica", 10, max_width)
                
                for line in lines:
                    if y_position < 50:
                        c.showPage()
                        y_position = height - 50
                    
                    c.drawString(margin, y_position, line)
                    y_position -= 12
                
                y_position -= 8
                    
            c.save()
            return pdf_path
            
        except Exception as e:
            logger.error(f"Erreur extraction DOCX: {e}")
            try:
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    pdf_path = tmp.name
                c = canvas.Canvas(pdf_path, pagesize=letter)
                c.setFont("Helvetica", 12)
                c.drawString(50, 750, f"Document: {Path(docx_path).name}")
                c.drawString(50, 730, "Erreur lors de l'extraction du contenu DOCX")
                c.save()
                return pdf_path
            except:
                return docx_path
    
    async def analyze_page_structure(self, image: Image.Image, page_num: int) -> Dict:
        """Analyse une page avec Claude Vision et extrait le texte complet"""
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        img_b64 = base64.b64encode(buffered.getvalue()).decode()
        analysis_prompt = """
        Analyse cette page de document médical et EXTRAIT TOUT LE TEXTE VISIBLE.

        Je veux deux choses :

        1. **TEXTE COMPLET** : Reproduis fidèlement TOUT le texte visible sur l'image, 
           en préservant la structure (titres, paragraphes, listes, tableaux).

        2. **STRUCTURE** : Identifie les sections logiques pour le découpage.

        Retourne un JSON structuré :
        ```json
        {
          "full_text": "TOUT LE TEXTE DE LA PAGE ICI, FORMATÉ AVEC DES RETOURS À LA LIGNE",
          "page_type": "guidelines|dosage_table|criteria_list",
          "sections": [
            {
              "title": "titre de la section",
              "type": "section|table|criteria|dosage|case_study",
              "text_content": "TEXTE COMPLET DE CETTE SECTION",
              "start_char": 0,
              "end_char": 150,
              "medical_entities": ["amoxicilline", "PAC grave"],
              "confidence": 0.9
            }
          ],
          "key_medical_info": {
            "medications": ["liste des médicaments"],
            "dosages": ["posologies identifiées"],
            "clinical_criteria": ["critères cliniques"],
            "patient_types": ["PAC grave", "sans comorbidité"]
          }
        }
        ```

        IMPORTANT : Le champ "full_text" doit contenir TOUT le texte de la page, 
        pas seulement un aperçu. Les sections doivent référencer des parties de ce texte complet.
        """
        
        try:
            import httpx
            import asyncio
            
            headers = {
                "Content-Type": "application/json",
                "x-api-key": self.client.api_key,
                "anthropic-version": "2023-06-01"
            }
            
            data = {
                "model": self.model,
                "max_tokens": 4000,  # Augmenté pour plus de texte
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/png",
                                    "data": img_b64
                                }
                            },
                            {
                                "type": "text",
                                "text": analysis_prompt
                            }
                        ]
                    }
                ]
            }
            
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    json=data,
                    headers=headers
                )
            
            if response.status_code == 200:
                result = response.json()
                response_text = result["content"][0]["text"]
                
                # Extraire le JSON de la réponse
                json_start = response_text.find('{')
                json_end = response_text.rfind('}') + 1
                
                if json_start != -1 and json_end != -1:
                    json_str = response_text[json_start:json_end]
                    analysis = json.loads(json_str)
                    analysis['page_number'] = page_num
                    return analysis
                else:
                    raise ValueError("No JSON found in response")
            else:
                raise ValueError(f"API error: {response.status_code}")
                
        except Exception as e:
            logger.error(f"Erreur analyse page {page_num}: {e}")
            return self._fallback_analysis(page_num)
    
    def _fallback_analysis(self, page_num: int) -> Dict:
        """Retourne une analyse de fallback en cas d'échec de l'API Claude"""
        return {
            "full_text": "",
            "page_type": "unknown",
            "page_number": page_num,
            "sections": [],
            "key_medical_info": {
                "medications": [],
                "dosages": [],
                "clinical_criteria": [],
                "patient_types": []
            }
        }

class IntelligentMedicalProcessor:
    """Processeur intelligent pour documents médicaux avec chunking adaptatif"""
    def __init__(self, anthropic_api_key: Optional[str] = None):
        self.analyzer = VisualDocumentAnalyzer(anthropic_api_key)
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            separators=[
                "\n\n*** ",
                "\n\n",
                "\n=== ",
                "\n• ",
                "\n- ",
                ". ",
                " "
            ]
        )
        
    async def process_medical_document(self, doc_path: str, progress_callback=None) -> List[LangChainDocument]:
        """Traite un document médical et retourne des chunks LangChain enrichis"""
        if progress_callback:
            await progress_callback(f"Conversion du document en images...", "vision")
        
        images = self.analyzer.convert_doc_to_images(doc_path)
        
        if progress_callback:
            await progress_callback(f"Document converti: {len(images)} pages à analyser", "vision")
        
        all_text_content = []
        page_analyses = []
        
        for i, image in enumerate(images):
            if progress_callback:
                await progress_callback(f"Extraction texte complet page {i+1}/{len(images)}...", "vision")
            
            analysis = await self.analyzer.analyze_page_structure(image, i)
            page_analyses.append(analysis)
            
            full_text = analysis.get('full_text', '')
            if full_text.strip():
                page_header = f"\n\n=== PAGE {i+1} ===\n"
                all_text_content.append(page_header + full_text)
            
            if progress_callback:
                text_length = len(full_text)
                sections_found = len(analysis.get('sections', []))
                await progress_callback(
                    f"Page {i+1}: {text_length} caractères, {sections_found} sections", 
                    "success"
                )
        
        complete_document_text = "\n".join(all_text_content)
        
        if progress_callback:
            await progress_callback(f"Texte total extrait: {len(complete_document_text)} caractères", "success")
            await progress_callback("Découpage intelligent en chunks...", "chunking")
        
        # Debug: Afficher le contenu total extrait
        logger.info("EXTRACTION DOCUMENTAIRE TERMINÉE:")
        logger.info("="*100)
        logger.info(f"   Nombre de pages analysées: {len(images)}")
        logger.info(f"   Texte total extrait: {len(complete_document_text)} caractères")
        logger.info(f"   Aperçu du texte (500 premiers chars): {complete_document_text[:500]}{'...' if len(complete_document_text) > 500 else ''}")
        logger.info("="*100)
        
        text_chunks = self.text_splitter.split_text(complete_document_text)
        
        logger.info("DÉCOUPAGE EN CHUNKS:")
        logger.info("="*100)
        logger.info(f"   Nombre de chunks créés: {len(text_chunks)}")
        logger.info(f"   Taille chunks: {settings.chunk_size} caractères")
        logger.info(f"   Chevauchement: {settings.chunk_overlap} caractères")
        logger.info("="*100)
        
        documents = []
        for chunk_idx, chunk_text in enumerate(text_chunks):
            page_num = self._find_page_for_chunk(chunk_text, page_analyses)
            medical_entities = self._extract_medical_entities_from_chunk(chunk_text, page_analyses)
            
            # Debug: Afficher chaque chunk créé
            logger.info(f"CHUNK VISION {chunk_idx+1}/{len(text_chunks)}:")
            logger.info(f"   Page source: {page_num}")
            logger.info(f"   Taille: {len(chunk_text)} caractères")
            logger.info(f"   Entités médicales: {medical_entities}")
            logger.info(f"   Contenu (200 premiers chars): {chunk_text[:200]}{'...' if len(chunk_text) > 200 else ''}")
            logger.info("   " + "-"*80)
            
            metadata = {
                'source': doc_path,
                'chunk_id': chunk_idx,
                'page': page_num,
                'chunk_size': len(chunk_text),
                'medical_entities': medical_entities,
                'document_type': 'medical_guidelines',
                'extraction_method': 'claude_vision_ocr',
                'total_chunks': len(text_chunks)
            }
            
            document = LangChainDocument(
                page_content=chunk_text,
                metadata=metadata
            )
            documents.append(document)
        
        logger.info(f"CRÉATION CHUNKS VISION TERMINÉE: {len(documents)} documents LangChain créés")
        logger.info("="*100)
        
        if progress_callback:
            await progress_callback(f"{len(documents)} chunks créés", "success")
        
        return documents
    
    def _find_page_for_chunk(self, chunk_text: str, page_analyses: List[Dict]) -> int:
        """Trouve la page source d'un chunk de texte"""
        for analysis in page_analyses:
            if analysis.get('full_text', '') in chunk_text or chunk_text in analysis.get('full_text', ''):
                return analysis.get('page_number', 0)
        return 0
    
    def _extract_medical_entities_from_chunk(self, chunk_text: str, page_analyses: List[Dict]) -> List[str]:
        """Extrait les entités médicales pertinentes pour un chunk donné"""
        entities = set()
        
        for analysis in page_analyses:
            medical_info = analysis.get('key_medical_info', {})
            for entity_type in ['medications', 'dosages', 'clinical_criteria', 'patient_types']:
                for entity in medical_info.get(entity_type, []):
                    if entity.lower() in chunk_text.lower():
                        entities.add(entity)
        
        return list(entities)

def create_medical_processor(anthropic_api_key: Optional[str] = None) -> IntelligentMedicalProcessor:
    """Factory pour créer un processeur médical intelligent"""
    return IntelligentMedicalProcessor(anthropic_api_key)